import os
import time
import streamlit as st
# from section_generator import generate_full_protocol
# from topic_name_sanitizer import sanitize_topic_name
# from cover_page_generator import generate_cover_page_from_protocol
from docx import Document

st.set_page_config(page_title="IRB Generator", layout="wide")

if "clear_inputs" not in st.session_state:
    st.session_state.clear()
    st.session_state.clear_inputs = True

st.markdown(
    """
    <h2 style='font-family: sans-serif; text-align: center;'>
        <span style='color: #F3A634;'>Tumerik\'s</span> IRB Generator
    </h2>
    
    <style>
    /* Style text input and textarea */
    .stTextInput > div > div > input,
    .stTextArea textarea {
        border: 2px solid #F3A634;
        background-color: #ffffff;
        color: #000000;
        padding: 8px;
        border-radius: 6px;
        font-size: 16px;
    }

    /* Optional: Adjust focus color */
    .stTextInput > div > div > input:focus,
    .stTextArea textarea:focus {
        border-color: #F3A634;
        box-shadow: 0 0 5px rgba(243, 166, 52, 0.5);
        outline: none;
    }
    </style>
    """,
    unsafe_allow_html=True)



# Initialize state
if "protocol_text" not in st.session_state:
    st.session_state.protocol_text = ""
if "submitted" not in st.session_state:
    st.session_state.submitted = False

# Inputs
study_title = st.text_input("Study Title", "")

st.markdown("### Study Summary Fields")

intervention = st.text_area("Intervention", height=100)
condition = st.text_area("Disease / Condition", height=100)
criteria = st.text_area("Selection / Exclusion Criteria", height=100)
objective = st.text_area("Primary Objective / Metric", height=100)
sponsor = st.text_area("Study Sponsor / Funding", height=100)

# Optional: combine into a dict or a full description string
study_summary = f"""
Intervention: {intervention}
Condition: {condition}
Selection/Exclusion Criteria: {criteria}
Primary Objective/Metric: {objective}
Sponsor/Funding: {sponsor}
"""

if st.button("Generate Full Protocol"):
    if not study_title or not study_summary.strip():
        st.warning("Please enter input into all fields.")
    else:
        with st.spinner("Thinking and drafting..."):
            time.sleep(5)  # Simulate delay
            # generate_full_protocol(study_title, study_summary)
            st.session_state.submitted = True

output_dir = "demo"
with open("demo/storm_gen_article_polished.txt", 'r') as file:
    text = file.read()
    st.session_state.protocol_text = text

# Display output
if st.session_state.submitted and st.session_state.protocol_text:
    st.subheader("Generated IRB Protocol")
    st.code(st.session_state.protocol_text, language="markdown")

    # Download as Word
    if st.button("Download Protocol as Word (.docx)"):
        # cover_text = generate_cover_page_from_protocol(st.session_state.protocol_text)

        doc = Document()
        # doc.add_heading("Protocol Summary", 0)
        # doc.add_paragraph(cover_text)

        doc.add_heading("Full Protocol", level=1)
        for paragraph in st.session_state.protocol_text.split("\n\n"):
            doc.add_paragraph(paragraph.strip())
        
        os.makedirs(output_dir, exist_ok=True)
        doc_path = os.path.join(output_dir, "IRB_Protocol.docx")
        doc.save(doc_path)

        with open(doc_path, "rb") as file:
            st.download_button(
                label="Download IRB_Protocol.docx",
                data=file,
                file_name="IRB_Protocol.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )